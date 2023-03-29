import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import Funcoes as f

def criar_documento(Dados, requisitos):
    documento = Document()
    style = documento.styles.add_style("Titulo", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    font.bold = True
    p = documento.add_paragraph()
    r = p.add_run()
    #r.add_picture('static/images/Completude.png', width=Inches(1.0), height=Inches(1.0))
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["Titulo"]
    run = paragrafo.add_run(
        "Relatório de Análise de Requisitos")
    paragrafo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    style = documento.styles.add_style("CORPO", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nSegue abaixo a análise de requisitos via ReqSCity, ferramenta essa desenvolvida num projeto PIBITI "
                            "apoiado pela Capes/CNPQ.")
    paragrafo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nRequisitos considerados ambíguos e X denota qual o motivo: ")
    tabela = documento.add_table(rows=len(Dados['Ambiguidade'])+1, cols=6)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = 'Table Grid'
    titulo_tabela = tabela.rows[0].cells
    titulos = ("Nº Requisito", "Palavra Ambigua", "Algoritmo Flexible Ambiguity", "Ambiguidade Analitica", "Ambiguidade por Coordenação", "Ambiguidade de ligação")
    
    for i in range(len(titulos)):
        titulo_tabela[i].text = titulos[i]
        titulo_tabela[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i in range(len(Dados['Ambiguidade'])):
        linha = tabela.rows[i+1].cells
        linha[0].text = str(Dados['Ambiguidade'][i][0])
        linha[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for j in range(len(Dados['Ambiguidade'][i][1])):
            if not Dados['Ambiguidade'][i][1][j]:
                linha[j+1].text = "\U00002713"
            else:
                linha[j + 1].text = "\U0001F5F4"
            linha[j+1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nRequisitos considerados incompletos e X denota qual o motivo: ")
    tabela = documento.add_table(rows=len(Dados['Analise_Sintatica'])+1, cols=5)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = 'Table Grid'
    titulo_tabela = tabela.rows[0].cells
    titulos = ("Nº Requisito", "Ausência de Verbos", "Voz Passiva", "Falta de Sujeito", "Dummy Subject")
    for i in range(len(titulos)) :
        titulo_tabela[i].text = titulos[i]
        titulo_tabela[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i in range(len(Dados['Analise_Sintatica'])):
        linha = tabela.rows[i+1].cells
        linha[0].text = str(Dados['Analise_Sintatica'][i][0])
        linha[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for j in range(len(Dados['Analise_Sintatica'][i][1])):
            if not Dados['Analise_Sintatica'][i][1][j]:
                linha[j+1].text = "\U00002713"
            else:
                linha[j + 1].text = "\U0001F5F4"
            linha[j+1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nRequisitos Considerados Bons: ")
    tabela = documento.add_table(rows=len(Dados['Bons'][0]) + 1, cols=1)
    tabela.style = 'Table Grid'
    titulo_tabela = tabela.rows[0].cells
    titulos = ["Nº do Requisito"]
    titulo_tabela[0].text=titulos[0]
    titulo_tabela[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i in range(len(Dados['Bons'][0])):
        linha = tabela.rows[i+1].cells
        linha[0].text = str(Dados['Bons'][0][i])
        linha[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nRequisitos Contextualizados: ")
    tabela = documento.add_table(rows=len(Dados['Contextualizados']) + 1, cols=1)
    tabela.style = 'Table Grid'
    titulo_tabela = tabela.rows[0].cells
    titulos = ["Nº do Requisito"]
    titulo_tabela[0].text=titulos[0]
    titulo_tabela[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i in range(len(Dados['Contextualizados'])):
        linha = tabela.rows[i+1].cells
        linha[0].text = str(Dados['Contextualizados'][i][0])
        linha[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nRequisitos Contextualizados porém com o sentido do sensor incompleto: ")
    checados = f.checa(Dados['Contextualizados'],1)
    tabela = documento.add_table(rows=len(checados) + 1, cols=1)
    tabela.style = 'Table Grid'
    titulo_tabela = tabela.rows[0].cells
    titulos = ["Nº do Requisito"]
    titulo_tabela[0].text=titulos[0]
    titulo_tabela[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    indices = []
    for i in range(len(Dados['Contextualizados'])):
        if Dados['Contextualizados'][i][1][1]:
            indices.append(Dados['Contextualizados'][i][0])
    for i in range(len(indices)):
        linha = tabela.rows[i+1].cells
        linha[0].text = str(indices[i])
        linha[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("\nRequisitos Contextualizados porém com o sentido do atuador incompleto: ")
    checados = f.checa(Dados['Contextualizados'],2)
    tabela = documento.add_table(rows=len(checados) + 1, cols=1)
    tabela.style = 'Table Grid'
    titulo_tabela = tabela.rows[0].cells
    titulos = ["Nº do Requisito"]
    titulo_tabela[0].text=titulos[0]
    titulo_tabela[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    indices = []
    for i in range(len(Dados['Contextualizados'])):
        if Dados['Contextualizados'][i][1][2]:
            indices.append(Dados['Contextualizados'][i][0])
    for i in range(len(indices)):
        linha = tabela.rows[i+1].cells
        linha[0].text = str(indices[i])
        linha[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    run = paragrafo.add_run("Requisitos que foram tratados: ")

    paragrafo = documento.add_paragraph()
    paragrafo.style = documento.styles["CORPO"]
    for texto in requisitos:
        run = paragrafo.add_run("{}\n".format(texto))
    documento.save("static\Documento.docx")